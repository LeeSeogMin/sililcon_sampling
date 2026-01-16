#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown to DOCX converter - Response document
Converts response.md to Word document with optimized table widths
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell


def set_cell_background(cell, fill_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_width(cell, width):
    """Set cell width"""
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width))
    tcW.set(qn('w:type'), 'auto')
    tcPr.append(tcW)


def set_cell_margins(cell, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set cell margins in mm (converted to twips: 1mm = 56.7 twips)"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for margin_name, margin_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin_elm = OxmlElement(f'w:{margin_name}')
        margin_elm.set(qn('w:w'), str(int(margin_val * 56.7)))  # mm to twips
        margin_elm.set(qn('w:type'), 'dxa')
        tcMar.append(margin_elm)

    tcPr.append(tcMar)


def parse_markdown(md_text):
    """Parse markdown into structured format"""
    lines = md_text.split('\n')
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Heading processing
        if line.startswith('# '):
            elements.append({'type': 'heading1', 'text': line[2:].strip()})
        elif line.startswith('## '):
            elements.append({'type': 'heading2', 'text': line[3:].strip()})
        elif line.startswith('### '):
            elements.append({'type': 'heading3', 'text': line[4:].strip()})

        # Quote processing
        elif line.startswith('> '):
            elements.append({'type': 'quote', 'text': line[2:].strip()})

        # Table processing
        elif line.startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            elements.append({'type': 'table', 'lines': table_lines})

        # Separator processing
        elif line.strip() == '---':
            elements.append({'type': 'separator'})

        # Code block processing
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append({'type': 'code', 'lines': code_lines})

        # Numbered list processing
        elif re.match(r'^\d+\.\s', line):
            elements.append({'type': 'list_item', 'text': re.sub(r'^\d+\.\s', '', line).strip()})

        # Regular text (ignore blank lines)
        elif line.strip():
            elements.append({'type': 'paragraph', 'text': line.strip()})

        i += 1

    return elements


def add_formatted_paragraph(doc, text, style=None):
    """Add paragraph with formatted text (bold, italic)"""
    p = doc.add_paragraph(style=style)

    # Split by formatting markers
    parts = re.split(r'(\*\*.*?\*\*|_.*?_|\[.*?\]\(.*?\))', text)

    for part in parts:
        if not part:
            continue
        elif part.startswith('**') and part.endswith('**'):
            # Bold text
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            # Italic text
            run = p.add_run(part[1:-1])
            run.italic = True
        elif re.match(r'\[.*?\]\(.*?\)', part):
            # Link
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                run = p.add_run(match.group(1))
                run.underline = True
        else:
            # Normal text
            if part:
                p.add_run(part)

    return p


def parse_table(table_lines):
    """Parse markdown table into rows and columns"""
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        # Skip separator rows (rows with dashes)
        if cells and not all(re.match(r'^-+$', cell) for cell in cells):
            rows.append(cells)
    return rows


def add_table_to_doc(doc, table_lines):
    """Add markdown table to document with optimized widths"""
    rows = parse_table(table_lines)

    if not rows:
        return

    # Separate headers from data
    headers = rows[0]
    data_rows = rows[1:]

    # Create table
    num_cols = len(headers)
    num_rows = len(data_rows) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Set autofit layout
    table.autofit = False
    table.allow_autofit = False

    # Calculate available width (landscape: 11 inches - 1 inch margins = 10 inches)
    available_width = Inches(10.0)
    col_width = available_width / num_cols

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.width = col_width

        # Set cell margins (1.0mm top/bottom/left/right)
        set_cell_margins(cell, top=1.0, bottom=1.0, left=1.0, right=1.0)

        # Clear default paragraph and add formatted text
        cell.paragraphs[0].text = ''
        add_formatted_paragraph(cell, header)

        # Format header
        set_cell_background(cell, 'D3D3D3')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.15  # 115% line spacing
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.width = col_width

            # Set cell margins (1.0mm top/bottom/left/right)
            set_cell_margins(cell, top=1.0, bottom=1.0, left=1.0, right=1.0)

            # Clear default paragraph and add formatted text
            cell.paragraphs[0].text = ''
            add_formatted_paragraph(cell, cell_text)

            # Format cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.line_spacing = 1.15  # 115% line spacing
                paragraph.word_wrap = True
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Set table width to 100% of page width
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Add table width element (5000 = 100% in pct type)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def create_word_document(md_file, output_file):
    """Create Word document from markdown file"""
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Parse markdown
    elements = parse_markdown(md_text)

    # Create Word document
    doc = Document()

    # Set page to landscape
    section = doc.sections[0]
    section.page_height = Inches(8.5)
    section.page_width = Inches(11.0)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Process elements
    for elem in elements:
        if elem['type'] == 'heading1':
            heading = doc.add_heading(elem['text'], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif elem['type'] == 'heading2':
            doc.add_heading(elem['text'], level=2)

        elif elem['type'] == 'heading3':
            doc.add_heading(elem['text'], level=3)

        elif elem['type'] == 'quote':
            p = doc.add_paragraph(elem['text'], style='Quote')
            p.paragraph_format.left_indent = Inches(0.25)

        elif elem['type'] == 'table':
            add_table_to_doc(doc, elem['lines'])
            doc.add_paragraph()  # Add spacing after table

        elif elem['type'] == 'separator':
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(6)

        elif elem['type'] == 'code':
            p = doc.add_paragraph()
            p.style = 'Normal'
            for line in elem['lines']:
                run = p.add_run(line + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)

        elif elem['type'] == 'list_item':
            doc.add_paragraph(elem['text'], style='List Number')

        elif elem['type'] == 'paragraph':
            add_formatted_paragraph(doc, elem['text'])

    # Save document
    doc.save(output_file)
    print(f"✓ Word document created: {output_file}")
    print(f"✓ Page orientation: Landscape (11\" x 8.5\")")
    print(f"✓ Table width: Optimized for page width")


if __name__ == '__main__':
    md_file = r'C:\Dev\silicon_sampling\docs\reply.md'
    output_file = r'C:\Dev\silicon_sampling\docs\reply.docx'

    create_word_document(md_file, output_file)
